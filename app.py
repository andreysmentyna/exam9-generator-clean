import os
import random
import zipfile
import tempfile
import subprocess
import json
import shutil
import logging
from pathlib import Path
from datetime import datetime, date

import requests
from flask import Flask, render_template, request, send_file, Response
from docx import Document
from docx.shared import Pt
from docxcompose.composer import Composer

# Настраиваем логирование, чтобы видеть этапы генерации
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.logger.setLevel(logging.INFO)

# Путь к папке с заданиями ( docx )
BANK_PATH = Path("bank")
GIST_ID = os.environ.get("GIST_ID")
GITHUB_TOKEN = os.environ.get("GITHUB_TOKEN")

if not GIST_ID or not GITHUB_TOKEN:
    raise RuntimeError("Set GIST_ID and GITHUB_TOKEN environment variables")

API_URL = f"https://api.github.com/gists/{GIST_ID}"
HEADERS = {
    "Authorization": f"token {GITHUB_TOKEN}",
    "Accept": "application/vnd.github.v3+json"
}

# ---------------------- Счётчик и статистика ----------------------
def get_counter_data():
    resp = requests.get(API_URL, headers=HEADERS)
    resp.raise_for_status()
    content = resp.json()["files"]["counter.json"]["content"]
    data = json.loads(content)
    if "history" not in data:
        data["history"] = {}
    if "value" not in data:
        data["value"] = 0
    return data

def update_counter_data(new_data):
    payload = {"files": {"counter.json": {"content": json.dumps(new_data)}}}
    requests.patch(API_URL, headers=HEADERS, json=payload).raise_for_status()

def get_global_counter(k):
    """
    Увеличивает общий счётчик на k и возвращает первый номер для этой пачки.
    Также записывает статистику по дням.
    """
    data = get_counter_data()
    current = data["value"]
    today = date.today().isoformat()
    data["value"] = current + k
    history = data["history"]
    history[today] = history.get(today, 0) + k
    update_counter_data(data)
    return current + 1

def get_stats():
    """
    Возвращает словарь со статистикой для отображения на главной странице:
    - today: количество сгенерированных сегодня
    - month: количество за текущий месяц
    - total: общее количество
    """
    data = get_counter_data()
    today = date.today().isoformat()
    today_count = data["history"].get(today, 0)
    current_month = today[:7]
    month_count = sum(
        count for day, count in data["history"].items()
        if day.startswith(current_month)
    )
    total = data["value"]
    return {"today": today_count, "month": month_count, "total": total}

# ---------- Оптимизированная конвертация docx → pdf ----------
def docx_to_pdf(docx_path, output_pdf_path):
    """
    Конвертирует DOCX в PDF с помощью LibreOffice.
    - Использует временный профиль в /dev/shm (виртуальный диск в памяти) для ускорения.
    - Устанавливает таймаут 600 секунд.
    """
    docx_path = Path(docx_path)
    output_pdf_path = Path(output_pdf_path)

    # Пытаемся создать временный профиль в RAM-диске для максимальной скорости.
    # Если не получается (например, /dev/shm отсутствует), используем обычный /tmp.
    try:
        tmp_profile = tempfile.mkdtemp(prefix="libre_", dir="/dev/shm")
        app.logger.info("LibreOffice будет использовать профиль в /dev/shm (RAM)")
    except Exception:
        tmp_profile = tempfile.mkdtemp(prefix="libre_")
        app.logger.info("Не удалось использовать /dev/shm, профиль в /tmp")

    env = os.environ.copy()
    env["HOME"] = str(tmp_profile)

    cmd = [
        "libreoffice",
        f"-env:UserInstallation=file://{tmp_profile / 'user'}",  # указываем профиль
        "--headless",
        "--nologo",
        "--norestore",
        "--invisible",
        "--convert-to", "pdf",
        "--outdir", str(output_pdf_path.parent),
        str(docx_path)
    ]

    try:
        app.logger.info(f"Запускаю конвертацию: {docx_path.name}")
        subprocess.run(cmd, check=True, timeout=600, env=env)
        expected_pdf = output_pdf_path.parent / (docx_path.stem + ".pdf")
        if expected_pdf.exists():
            os.rename(expected_pdf, output_pdf_path)
            app.logger.info(f"Успешно сконвертирован в {output_pdf_path.name}")
        else:
            raise FileNotFoundError("LibreOffice не создал ожидаемый PDF-файл")
    except Exception as e:
        app.logger.error(f"Ошибка конвертации LibreOffice: {e}")
        raise RuntimeError(f"Ошибка конвертации LibreOffice: {e}")
    finally:
        # Удаляем временный профиль, чтобы не засорять память
        shutil.rmtree(tmp_profile, ignore_errors=True)

# ---------- Умные заголовки вариантов ----------
def determine_header(selected_tasks):
    """Возвращает заголовок в зависимости от выбранных заданий."""
    if not selected_tasks:
        return "Вариант №"
    if set(selected_tasks) == set(range(1, 19)):
        return "Вариант экзаменационной работы №"
    if set(selected_tasks) == set(range(1, 14)):
        return "Вариант № (часть 1, задания 1–13)"
    if set(selected_tasks) == set(range(14, 19)):
        return "Вариант № (часть 2, задания 14–18)"
    if set(selected_tasks) == set(list(range(1, 10)) + list(range(14, 17))):
        return "Вариант № (Алгебра)"
    if set(selected_tasks) == set(list(range(10, 14)) + list(range(17, 19))):
        return "Вариант № (Геометрия)"
    # Произвольный набор
    nums = sorted(selected_tasks)
    return f"Вариант № (задания {', '.join(map(str, nums))})"

def add_title_paragraph(master, text):
    """Добавляет в документ жирный центрированный параграф с заголовком."""
    para = master.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return para

# ---------- Основная логика генерации ----------
def compose_variants(k, selected_tasks, preview=False, merge=False):
    """
    Собирает k вариантов, используя задания из списка selected_tasks.
    :param k: количество вариантов
    :param selected_tasks: список номеров заданий (1..18)
    :param preview: если True, счётчик не увеличивается
    :param merge: если True, все варианты объединяются в один PDF
    :return: кортеж (путь к итоговому файлу, стартовый глобальный номер)
    """
    # Собираем список папок с заданиями, отфильтрованных по выбранным номерам
    task_dirs = sorted(
        [d for d in BANK_PATH.iterdir() if d.is_dir() and d.name.startswith("task_")],
        key=lambda x: int(x.name.split("_")[1])
    )
    filtered_dirs = [d for d in task_dirs if int(d.name.split("_")[1]) in selected_tasks]
    if not filtered_dirs:
        raise Exception("Не выбрано ни одного задания")

    # Для каждого задания загружаем список файлов вариантов и перемешиваем
    all_tasks = []
    for task_dir in filtered_dirs:
        docx_files = sorted(
            [f for f in task_dir.glob("variant_*.docx")],
            key=lambda x: int(x.stem.split("_")[1])
        )
        if not docx_files:
            raise Exception(f"В папке {task_dir.name} нет docx-файлов")
        random.shuffle(docx_files)
        all_tasks.append(docx_files)

    # Получаем глобальный номер (если не превью)
    if not preview:
        start_global_id = get_global_counter(k)
    else:
        start_global_id = 0

    # Создаём временную папку для сборки
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        result_files = []   # здесь будут либо docx (для merge), либо pdf

        for exam_idx in range(k):
            global_id = start_global_id + exam_idx
            master = Document()
            composer = Composer(master)

            # Добавляем заголовок варианта (без отрыва страницы)
            title = f"{determine_header(selected_tasks)} {global_id}"
            if preview:
                title += " (ПРИМЕР)"
            add_title_paragraph(master, title)

            # Добавляем все выбранные задания
            for variants_list in all_tasks:
                variant_file = variants_list[exam_idx % len(variants_list)]
                sub_doc = Document(str(variant_file))
                composer.append(sub_doc)

            # Сохраняем собранный docx
            docx_path = tmpdir_path / f"variant_{global_id}.docx"
            master.save(str(docx_path))

            if merge:
                # В режиме merge откладываем docx для последующего объединения
                result_files.append(docx_path)
            else:
                # Сразу конвертируем в PDF
                pdf_path = tmpdir_path / f"variant_{global_id}.pdf"
                docx_to_pdf(docx_path, pdf_path)
                result_files.append(pdf_path)

        # Финальная сборка результата
        if merge:
            # Объединяем все docx в один файл и конвертируем в PDF
            merged_doc = Document()
            merged_composer = Composer(merged_doc)
            first = True
            for doc_path in result_files:
                if not first:
                    merged_doc.add_page_break()   # разрыв страницы между вариантами
                merged_composer.append(Document(str(doc_path)))
                first = False
            merged_docx = tmpdir_path / "merged.docx"
            merged_doc.save(str(merged_docx))
            merged_pdf = tmpdir_path / "merged.pdf"
            docx_to_pdf(merged_docx, merged_pdf)
            # Перемещаем финальный PDF в постоянную временную папку
            final_path = Path(tempfile.gettempdir()) / f"variants_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
            merged_pdf.rename(final_path)
            return final_path, start_global_id
        else:
            # Пакуем все PDF в ZIP
            zip_path = tmpdir_path / "variants.zip"
            with zipfile.ZipFile(zip_path, "w") as zf:
                for file_path in result_files:
                    zf.write(file_path, arcname=file_path.name)
            final_zip = Path(tempfile.gettempdir()) / f"variants_{datetime.now().strftime('%Y%m%d%H%M%S')}.zip"
            zip_path.rename(final_zip)
            return final_zip, start_global_id

# ---------------------- Маршруты Flask ----------------------
@app.route("/")
def index():
    """Главная страница с формой и статистикой."""
    stats = {}
    try:
        stats = get_stats()
    except Exception as e:
        app.logger.error(f"Ошибка получения статистики: {e}")
        stats = {"today": "—", "month": "—", "total": "—"}
    return render_template("index.html", stats=stats)

@app.route("/generate", methods=["POST"])
def generate():
    """Обрабатывает запрос на генерацию вариантов."""
    try:
        k = int(request.form.get("quantity", 0))
        merge = request.form.get("merge") == "1"

        if k < 1 or k > 50:
            return "Количество вариантов должно быть от 1 до 50", 400

        # Для большого числа вариантов принудительно включаем объединение,
        # чтобы избежать множества отдельных конвертаций и возможного таймаута.
        if k > 20 and not merge:
            merge = True
            app.logger.info(f"Автоматически включено объединение для {k} вариантов")

        selected_tasks = [int(t) for t in request.form.getlist("task")]
        if not selected_tasks:
            return "Не выбрано ни одного задания", 400
        if not all(1 <= t <= 18 for t in selected_tasks):
            return "Некорректный номер задания", 400

        result_file, start_id = compose_variants(k, selected_tasks, preview=False, merge=merge)

        # Формируем понятное имя файла
        if k == 1:
            base_name = f"exam-var-{start_id}"
        else:
            base_name = f"exam-vars-{start_id}-{start_id + k - 1}"

        if merge:
            download_name = f"{base_name}.pdf"
            return send_file(result_file, as_attachment=True,
                             download_name=download_name, mimetype="application/pdf")
        else:
            download_name = f"{base_name}.zip"
            return send_file(result_file, as_attachment=True,
                             download_name=download_name)
    except Exception as e:
        app.logger.error(f"Ошибка генерации: {e}")
        return f"Ошибка при генерации: {e}", 500

@app.route("/preview", methods=["POST"])
def preview():
    """Генерирует один вариант без увеличения счётчика и показывает его в браузере."""
    try:
        selected_tasks = [int(t) for t in request.form.getlist("task")]
        if not selected_tasks:
            return "Не выбрано ни одного задания", 400
        if not all(1 <= t <= 18 for t in selected_tasks):
            return "Некорректный номер задания", 400

        result_file, _ = compose_variants(1, selected_tasks, preview=True, merge=False)
        with zipfile.ZipFile(result_file, "r") as zf:
            pdf_name = zf.namelist()[0]
            pdf_data = zf.read(pdf_name)
        return Response(pdf_data, mimetype="application/pdf",
                        headers={"Content-Disposition": "inline; filename=preview.pdf"})
    except Exception as e:
        app.logger.error(f"Ошибка предпросмотра: {e}")
        return f"Ошибка при предпросмотре: {e}", 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000, debug=False)